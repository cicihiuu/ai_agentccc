from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "交付"


def make_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return doc


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for idx, header in enumerate(headers):
        t.rows[0].cells[idx].text = header
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build_01() -> None:
    doc = make_doc("项目现状与已完成工作说明", "面向接手开发者的当前交付状态说明")

    h(doc, "一、项目目标与当前基线")
    bullets(
        doc,
        [
            "本项目的目标来自《基于AI Agent的网络安全工具开发》，目标是构建面向授权目标 URL 的 AI Agent 安全测试工具。",
            "项目能力范围不局限于单一靶场，理论上面向授权课程靶场、本地 Docker 靶场以及其他授权测试目标。",
            "当前这次交付和验收默认基线是 Pikachu Docker，地址为 http://127.0.0.1:8765/。",
            "当前工作台地址为 http://127.0.0.1:8000/，默认解释器为 D:\\Python\\python.exe。",
        ],
    )

    h(doc, "二、当前已经完成的内容")
    numbers(
        doc,
        [
            "完成统一接口：Finding、ModuleResult、ScanRun，可统一承接各模块结果。",
            "完成 AgentRuntime 四阶段：plan、step、step_replan、final_answer。",
            "完成 Human-in-the-loop：sql_scan、poc_verify 作为高风险步骤必须人工确认。",
            "完成状态持久化：运行过程写入 python_mvp/runs/，可以恢复 run。",
            "完成 Web/API 工作台：可创建任务、继续执行、查看状态、批准/拒绝高风险步骤、生成 HTML 报告。",
            "完成 Ollama 适配与规则回退。",
            "完成白名单内的 recon、backup_audit、js_audit、sql_scan、poc_verify 五个模块的当前版本。",
            "完成当前测试基线：python -m unittest discover -s python_mvp/tests 为 30 tests OK。",
        ],
    )

    h(doc, "三、当前还没有完成的内容")
    bullets(
        doc,
        [
            "参考书要求的 12+ Skill 还没有补齐，目前只完成 5 个主 skill。",
            "SQL/WAF bypass 还没有达到参考书中的完整目标，当前仍处于受控候选参数识别和验证链路阶段。",
            "JS AST 深化分析、误报控制、行号级定位仍需增强。",
            "备份文件下载、解压、静态分析、源码深审仍需增强。",
            "PDF 报告尚未完成。",
            "多靶场与更通用的环境兼容说明仍需完善。",
        ],
    )

    h(doc, "四、当前交互版和旧 MVP 的差异")
    bullets(
        doc,
        [
            "旧 MVP 主要面向 fixture、静态演示报告和统一接口。",
            "当前交互版新增了 API、工作台、运行状态、事件日志、人工批准和更真实的靶场探测。",
            "旧材料很多使用 DVWA:8080 作为示例；当前交付基线已经切换到 Pikachu:8765。",
            "交付口径必须区分“项目目标不局限单一靶场”和“当前默认验收环境为 Pikachu:8765”。",
        ],
    )

    h(doc, "五、接手者的第一优先级")
    bullets(
        doc,
        [
            "先理解当前代码结构和工作台行为，不要急着扩功能。",
            "先在 Pikachu:8765 环境下复现当前验收链路。",
            "再按 B/C/D 的详细任务书分别推进 SQL/POC、前端 JS、Recon/Backup/环境维护。",
        ],
    )

    doc.save(OUT_DIR / "01_我的任务清单_周报版.docx")


def build_02() -> None:
    doc = make_doc("交付说明（详细版）", "项目总说明 + 当前代码说明 + 当前环境说明 + 接手开发说明")

    h(doc, "一、项目整体要做什么")
    para(
        doc,
        "本项目来自《基于AI Agent的网络安全工具开发》。参考书要求的不是某一个固定靶场脚本，"
        "而是一套面向授权目标 URL 的 AI Agent 网络安全工具。它的理想目标包括：多 Agent 架构、Profile、Skill、"
        "SQL 检测与验证、前端 JS 审计、备份文件审计、POC 验证、LLM 参与、报告生成和完整验收链路。"
    )
    bullets(
        doc,
        [
            "项目不绑定单一靶场。",
            "项目目标/能力范围：不绑定单一靶场，面向授权目标 URL 与课程靶场。",
            "当前默认验收环境为 Pikachu:8765。",
            "当前默认验收环境：Pikachu:8765 + Workbench:8000 + D:\\Python\\python.exe。",
            "当前交付环境只是本次默认复现基线，不等于项目永久只能跑 Pikachu。",
        ],
    )

    h(doc, "二、当前已完成什么")
    bullets(
        doc,
        [
            "统一接口：Finding / ModuleResult / ScanRun 已完成。",
            "AgentRuntime 四阶段：plan / step / step_replan / final_answer 已接通。",
            "Web/API 工作台已完成：创建任务、执行、暂停、批准、拒绝、报告。",
            "Ollama 适配与规则回退已完成。",
            "基于白名单的 recon / backup / js / sql / poc 模块已完成当前版本。",
            "状态持久化、事件流和人工确认链已完成。",
            "当前测试已通过：30 tests OK。",
        ],
    )

    h(doc, "三、当前还需要进一步完善什么")
    bullets(
        doc,
        [
            "12+ Skill 尚未补齐。",
            "SQL/WAF bypass 尚未达到参考书的完整要求。",
            "JS AST 深化、误报控制、定位能力需要进一步增强。",
            "备份文件下载、解压、静态分析尚未接通。",
            "PDF 报告尚未完成。",
            "多靶场兼容说明和环境迁移说明仍需增强。",
        ],
    )

    h(doc, "四、当前项目结构是什么")
    table(
        doc,
        ["目录/模块", "作用", "当前状态", "后续如何利用"],
        [
            ["agent/", "运行时主链、计划、状态、子任务、门禁", "已接通", "后续所有能力扩展都要沿用这里"],
            ["api/", "FastAPI 工作台和服务层", "已接通", "当前正式验收主要通过这里进行"],
            ["llm/", "LLM provider 层", "已接通", "后续切换模型或增强规划逻辑时继续沿用"],
            ["modules/", "安全能力实现", "已接通但不完整", "B/C/D 的主要开发区域"],
            ["profiles/", "Profile 定义", "已接通", "后续新增角色或模式时继续补充"],
            ["skills/", "Skill 元数据", "仅 5 个", "后续扩展到 12+ 时继续使用"],
            ["tests/", "单元测试", "已通过", "后续所有开发都要同步维护"],
            ["scripts/", "环境与文档脚本", "已存在", "用于部署、恢复、交付文档再生成"],
            ["runs/", "运行状态与结果快照", "已可写入", "用于恢复 run，不作为手工编辑目标"],
            ["reports/", "HTML 报告输出", "可生成", "后续可继续扩展到 PDF"],
        ],
    )

    h(doc, "五、python_mvp 全量文件级说明")
    h(doc, "5.1 根目录文件", level=2)
    table(
        doc,
        ["文件", "作用", "当前状态", "后续是否还要改", "通常由谁改"],
        [
            ["README.md", "人类阅读入口，说明如何运行和验收", "已同步到 Pikachu 基线", "需要继续维护", "组长"],
            ["requirements.txt", "工作台依赖", "包含 fastapi/uvicorn", "需要继续维护", "组长"],
            ["run_demo.py", "CLI 演示入口", "仍可用", "可保留", "组长"],
            ["fixtures/demo_run.json", "旧静态 fixture", "保留作旧链路示例", "一般不重点修改", "组长"],
        ],
    )

    h(doc, "5.2 profiles", level=2)
    table(
        doc,
        ["文件", "作用", "当前状态", "后续是否还要改", "通常由谁改"],
        [
            ["blackbox_pentest.yaml", "黑盒主流程 profile", "可用", "可能继续扩展", "组长/B"],
            ["frontend_audit.yaml", "前端审计 profile", "可用", "可能继续扩展", "组长/C"],
        ],
    )

    h(doc, "5.3 skills", level=2)
    table(
        doc,
        ["文件", "作用", "当前状态", "后续是否还要改", "通常由谁改"],
        [
            ["recon.yaml", "recon skill 说明", "可用", "可能继续扩展", "D/组长"],
            ["backup_audit.yaml", "backup skill 说明", "可用", "可能继续扩展", "D/组长"],
            ["js_audit.yaml", "JS 审计 skill 说明", "可用", "需要增强", "C/组长"],
            ["sql_scan.yaml", "SQL skill 说明", "可用", "需要增强", "B/组长"],
            ["poc_verify.yaml", "POC skill 说明", "可用", "需要增强", "B/组长"],
        ],
    )

    h(doc, "5.4 src/ai_security_agent", level=2)
    table(
        doc,
        ["文件", "作用", "当前状态", "后续是否还要改", "通常由谁改"],
        [
            ["schemas.py", "统一接口定义", "稳定核心", "尽量不改", "组长"],
            ["main.py", "CLI 主入口", "默认目标已指向 8765", "必要时可调", "组长"],
            ["orchestrator.py", "旧总控与汇总逻辑", "CLI 仍在使用", "一般少改", "组长"],
            ["report.py", "HTML 报告生成", "当前可用", "需要继续维护", "组长"],
            ["api/app.py", "FastAPI 路由层", "当前可用", "需要继续维护", "组长"],
            ["api/service.py", "工作台服务层", "当前可用", "需要继续维护", "组长"],
            ["api/static/index.html", "工作台前端页面", "当前可用", "需要继续维护", "组长"],
            ["llm/base.py", "LLM 抽象", "当前可用", "需要继续维护", "组长"],
            ["llm/ollama.py", "Ollama provider", "当前可用", "需要继续维护", "组长"],
            ["agent/profile_loader.py", "Profile 读取", "当前可用", "一般少改", "组长"],
            ["agent/skill_registry.py", "Skill 注册与筛选", "当前可用", "一般少改", "组长"],
            ["agent/planner.py", "规则/LLM 规划", "当前可用", "后续可能增强", "组长"],
            ["agent/runtime.py", "Agent 主链", "当前可用", "谨慎改", "组长"],
            ["agent/human_gate.py", "人工确认门禁", "当前可用", "谨慎改", "组长"],
            ["agent/state_store.py", "状态持久化", "当前可用", "谨慎改", "组长"],
            ["agent/subagent.py", "子任务调度", "当前可用", "谨慎改", "组长"],
            ["modules/common.py", "通用抓取与白名单", "当前可用", "可能增强", "组长/D"],
            ["modules/recon.py", "首页与基础入口识别", "当前可用", "后续继续增强", "D"],
            ["modules/backup_audit.py", "备份与敏感文件探测", "当前可用", "后续继续增强", "D"],
            ["modules/js_audit.py", "前端 JS 审计", "当前可用", "后续继续增强", "C"],
            ["modules/sql_scan.py", "SQL 候选参数识别", "当前可用", "后续继续增强", "B"],
            ["modules/poc_verify.py", "POC 验证记录", "当前可用", "后续继续增强", "B"],
        ],
    )

    h(doc, "5.5 tests", level=2)
    table(
        doc,
        ["文件", "作用", "当前状态", "后续是否还要改", "通常由谁改"],
        [
            ["test_schemas.py", "接口约束测试", "稳定", "一般少改", "组长"],
            ["test_orchestrator.py", "旧总控测试", "稳定", "一般少改", "组长"],
            ["test_main.py", "CLI 报告入口测试", "稳定", "一般少改", "组长"],
            ["test_report.py", "报告生成测试", "稳定", "一般少改", "组长"],
            ["test_modules.py", "模块基础行为测试", "重要", "会跟着模块改", "B/C/D"],
            ["test_agent_profile.py", "Profile 读取测试", "稳定", "一般少改", "组长"],
            ["test_agent_runtime.py", "Runtime 测试", "重要", "可能会改", "组长/B/C"],
            ["test_agent_stage_5_6_7.py", "高风险门禁与恢复测试", "重要", "可能会改", "组长/B"],
            ["test_skill_registry.py", "Skill 注册测试", "稳定", "一般少改", "组长"],
        ],
    )

    h(doc, "5.6 scripts", level=2)
    table(
        doc,
        ["文件", "作用", "当前状态", "后续是否还要改", "通常由谁改"],
        [
            ["start_pikachu.ps1", "Pikachu 启动与恢复脚本", "当前可用", "需要维护", "D/组长"],
            ["start_workbench.ps1", "Workbench 启动脚本", "当前可用", "需要维护", "组长/D"],
            ["generate_delivery_docs_v3.py", "整套交付文档生成脚本", "当前新增", "后续文档变化时继续改", "组长"],
        ],
    )

    h(doc, "六、网页到底在做什么")
    bullets(
        doc,
        [
            "GET /：返回工作台页面，页面显示 profile、目标、计划、状态、事件和报告入口。",
            "GET /api/profiles：返回当前可选 profile 列表。",
            "POST /api/runs：创建新的 run，后端会加载 profile、注册 skill、生成 execution plan、保存 state。",
            "POST /api/runs/{run_id}/step：执行一步或一个 ready batch。",
            "POST /api/runs/{run_id}/continue：一直执行到完成或遇到高风险确认点。",
            "POST /api/runs/{run_id}/approve/{step_id}：批准某个高风险步骤。",
            "POST /api/runs/{run_id}/deny/{step_id}：拒绝某个高风险步骤。",
            "POST /api/runs/{run_id}/report：把当前 run 写成 HTML 报告。",
            "页面上的 Start、Continue、Approve、Deny、Generate Report 按钮都与这些后端动作一一对应。",
        ],
    )

    h(doc, "七、每个阶段产出什么 py、做什么、现有怎么利用")
    table(
        doc,
        ["阶段", "核心产物", "作用", "当前如何利用"],
        [
            ["Phase 0-1", "schemas.py, demo_run.json, test_schemas.py", "统一接口和静态示例", "作为后续模块和报告的统一数据入口"],
            ["Phase 2-3", "orchestrator.py, main.py, report.py, test_report.py", "总控与报告生成", "CLI 和旧链路仍在使用"],
            ["Phase 4-7", "modules/*.py, tests/*, final_demo_report.html", "模块接入、联调、演示和阶段测试", "当前交互版继续沿用这些模块结构"],
            ["当前交互版新增", "api/*.py, llm/*.py, agent/runtime.py, scripts/*.ps1", "工作台、LLM、事件流、恢复、环境脚本", "当前正式验收和交付主要依赖这部分"],
        ],
    )

    h(doc, "八、当前交付环境")
    bullets(
        doc,
        [
            "当前默认验收环境是 Pikachu:8765。",
            "这是本次交付基线，不等于项目永久绑定 Pikachu。",
            "项目本身面向授权目标 URL 与课程靶场，后续可以切到 DVWA、bWAPP 或其他授权靶场。",
            "如果切换靶场，通常需要同步检查默认 target、测试示例、脚本、截图、模块 evidence 说明。",
        ],
    )

    h(doc, "九、部署与修改提示")
    bullets(
        doc,
        [
            "当前解释器基线是 D:\\Python\\python.exe，不建议混用 Inkscape 自带 python。",
            "PowerShell 默认可能拦截 .ps1，必要时使用 -ExecutionPolicy Bypass 或手工执行 docker 命令。",
            "Pikachu 若出现 8765 空响应，通常需要删除 apache pid 后重启 apache2。",
            "如果别人机器路径不同，通常要改 README、交付说明和启动脚本中的路径示例。",
            "如果别人解释器不是 D:\\Python\\python.exe，需要统一改成他们自己的 python，并重新验证 fastapi、uvicorn、python-docx 的安装位置。",
        ],
    )

    h(doc, "十、历史文档如何使用")
    bullets(
        doc,
        [
            "docs/组长 和 docs/A 下的大部分阶段文档用于记录开发过程和早期材料。",
            "历史文档里可能保留 DVWA:8080 或旧阶段口径，这些材料可以用于追踪来路，但不应再作为当前交付标准。",
            "若历史阶段文档与本交付说明冲突，以本交付说明为准。",
        ],
    )

    doc.save(OUT_DIR / "02_交付说明_环境与部署.docx")


def build_03() -> None:
    doc = make_doc("BCD 任务索引总览", "详细任务书导航、边界总则与联调顺序")
    h(doc, "一、三人负责方向")
    table(
        doc,
        ["成员", "方向", "主要文件"],
        [
            ["B", "SQL 与 POC", "sql_scan.py / poc_verify.py / 对应测试"],
            ["C", "前端 JS 审计", "js_audit.py / 对应测试"],
            ["D", "Recon、Backup、靶场环境与证据", "recon.py / backup_audit.py / scripts / 交付说明"],
        ],
    )
    h(doc, "二、三人统一遵守的规则")
    bullets(
        doc,
        [
            "统一接口不改：Finding、ModuleResult、ScanRun。",
            "高风险门禁不绕过：sql_scan、poc_verify 继续保留人工批准。",
            "项目不局限单一靶场。",
            "当前默认验收环境是 Pikachu:8765，但项目目标不局限单一靶场。",
            "不要把项目做成公网扫描器或自动攻击器。",
        ],
    )
    h(doc, "三、详细任务书文件")
    bullets(
        doc,
        [
            "04_B_task_handoff.docx：B 的详细任务说明",
            "05_C_task_handoff.docx：C 的详细任务说明",
            "06_D_task_handoff.docx：D 的详细任务说明",
        ],
    )
    h(doc, "四、联调顺序")
    numbers(
        doc,
        [
            "先确保 Pikachu:8765 和 Workbench:8000 可启动。",
            "再确认 frontend_audit 跑通。",
            "然后推进 blackbox_pentest 的 sql_scan 人工确认链路。",
            "最后整理报告、截图和交付说明。",
        ],
    )
    doc.save(OUT_DIR / "03_BCD任务清单_详细版.docx")


def build_member_doc(
    out_name: str,
    title: str,
    subtitle: str,
    role_intro: str,
    status_items: list[str],
    read_items: list[str],
    file_rows: list[list[str]],
    no_change_items: list[str],
    phases: list[tuple[str, list[str], str]],
    accept_steps: list[str],
    expected_items: list[str],
    deliver_items: list[str],
    report_text: str,
) -> None:
    doc = make_doc(title, subtitle)
    h(doc, "一、当前角色定位")
    para(doc, role_intro)

    h(doc, "二、必须先知道的项目现状")
    bullets(doc, status_items)

    h(doc, "三、需要重点阅读的文件")
    numbers(doc, read_items)

    h(doc, "四、负责哪些文件")
    table(doc, ["文件", "负责什么"], file_rows)

    h(doc, "五、不要改哪些地方")
    bullets(doc, no_change_items)

    h(doc, "六、阶段任务")
    for phase_title, phase_items, phase_done in phases:
        h(doc, phase_title, level=2)
        bullets(doc, phase_items)
        para(doc, phase_done)

    h(doc, "七、验收步骤")
    numbers(doc, accept_steps)

    h(doc, "八、做完后应该看到什么结果")
    bullets(doc, expected_items)

    h(doc, "九、最终交付物")
    bullets(doc, deliver_items)

    h(doc, "十、给组长汇报时的表达模板")
    para(doc, report_text)

    doc.save(OUT_DIR / out_name)


def build_04() -> None:
    build_member_doc(
        "04_B_task_handoff.docx",
        "成员 B 详细任务说明",
        "SQL 检测与 POC 验证 / 基于当前 Pikachu 8765 交付基线",
        "B 的职责不是重写整个 Agent 系统，也不是去做前端 JS 审计、备份审计、工作台页面或 LLM 规划。"
        "B 的核心任务，是基于当前已经跑通的 Pikachu 8765 交付基线，把 sql_scan 和 poc_verify 从受控演示版推进到更真实、"
        "更有证据链、适合答辩展示的靶场验证版，同时不破坏统一接口、人工确认门禁和现有报告结构。",
        [
            "项目目标不局限单一靶场。",
            "当前交付与验收默认基线是 Pikachu:8765。",
            "当前默认靶场是 Docker Pikachu，地址是 http://127.0.0.1:8765/。",
            "当前工作台地址是 http://127.0.0.1:8000/，支持创建任务、继续执行、人工批准高风险步骤、生成报告。",
            "当前 Agent 主链位于 python_mvp/src/ai_security_agent/agent/runtime.py，sql_scan 和 poc_verify 都被定义为高风险步骤。",
            "当前项目目标是授权靶场验证，不是公网扫描器，更不是自动攻击器。",
        ],
        [
            "python_mvp/src/ai_security_agent/schemas.py",
            "python_mvp/src/ai_security_agent/modules/sql_scan.py",
            "python_mvp/src/ai_security_agent/modules/poc_verify.py",
            "python_mvp/src/ai_security_agent/agent/runtime.py",
            "python_mvp/src/ai_security_agent/api/service.py",
            "python_mvp/tests/test_modules.py、test_agent_stage_5_6_7.py",
            "python_mvp/skills/sql_scan.yaml 和 python_mvp/skills/poc_verify.yaml",
        ],
        [
            [r"python_mvp/src/ai_security_agent/modules/sql_scan.py", "SQL 候选参数识别、页面定位、证据链表达"],
            [r"python_mvp/src/ai_security_agent/modules/poc_verify.py", "把 POC 结果改造成更完整的受控验证记录"],
            [r"python_mvp/skills/sql_scan.yaml", "维护 SQL skill 的说明、触发词、检查清单"],
            [r"python_mvp/skills/poc_verify.yaml", "维护 POC skill 的说明、触发词、检查清单"],
            [r"python_mvp/tests/test_modules.py", "补充与 SQL/POC 实际行为匹配的模块测试"],
        ],
        [
            "不要修改 schemas.py 的字段结构。",
            "不要重写 api/app.py、api/service.py、agent/runtime.py 的主链逻辑，除非组长明确要求。",
            "不要加入公网扫描、自动攻击、危险 payload 自动生成等超范围能力。",
            "不要让 sql_scan 或 poc_verify 绕过人工确认直接执行高风险动作。",
        ],
        [
            (
                "阶段 0：理解边界",
                [
                    "确认当前靶场是 Pikachu 8765 本地 Docker 环境。",
                    "确认 SQL 和 POC 依然属于高风险步骤，必须经过人工批准。",
                    "确认自己的输出目标是更真实的靶场验证记录，而不是更强的攻击能力。",
                ],
                "阶段 0 通过标准：能解释为什么 sql_scan 和 poc_verify 必须保留人工确认门禁。",
            ),
            (
                "阶段 1：完善 sql_scan",
                [
                    "尽量写明页面、参数名、提取来源和判断依据。",
                    "优先利用 Pikachu 页面里的 query 参数、form input 名称、提交入口、菜单路径构建 evidence。",
                    "如果拿不到参数，也要把为什么未确认写成清楚 evidence，而不是只有含糊日志。",
                ],
                "阶段 1 通过标准：批准 sql_scan 后，报告里的 SQL finding 至少要能说明页面、参数和判断依据。",
            ),
            (
                "阶段 2：完善 poc_verify",
                [
                    "让 poc_verify 基于 high_findings 生成更完整的验证记录，而不是一句占位说明。",
                    "结果至少包含原始 finding 标题、位置、验证方式、验证结论、后续建议。",
                    "没有 high finding 或没有人工批准时，必须继续保持 skipped。",
                ],
                "阶段 2 通过标准：批准 sql_scan 和 poc_verify 后，报告中出现结构完整的 POC verification record。",
            ),
        ],
        [
            r"cd D:\App\aster-main\aster-main",
            r"python -m unittest discover -s python_mvp/tests",
            r"用 D:\Python\python.exe 启动工作台，打开 http://127.0.0.1:8000/",
            r"目标填 http://127.0.0.1:8765/ ，Profile 选 blackbox_pentest",
            r"点击 Continue Until Pause，确认 sql_scan 停在待批准",
            r"批准 sql_scan 后继续执行，再批准 poc_verify，最后生成报告",
        ],
        [
            "第一次 Continue Until Pause：recon、backup_audit、js_audit 完成，sql_scan 等待批准。",
            "批准 sql_scan 后：sql_scan 变为 ok，并写出更具体的参数名、页面路径或判断依据。",
            "批准 poc_verify 后：报告里出现一条更完整的 POC verification record。",
        ],
        [
            "更新后的 sql_scan.py",
            "更新后的 poc_verify.py",
            "同步过的 sql_scan.yaml 与 poc_verify.yaml",
            "配套测试更新",
            "SQL 等待确认截图、批准后截图、最终报告截图",
        ],
        "我没有改 Agent 主链，也没有改统一接口。我主要增强了 sql_scan 和 poc_verify 的证据链，让它们在 Pikachu 8765 靶场中更像真实验证记录。",
    )


def build_05() -> None:
    build_member_doc(
        "05_C_task_handoff.docx",
        "成员 C 详细任务说明",
        "前端 JS 审计 / 基于当前 Pikachu 8765 交付基线",
        "C 的职责不是重写 Agent 主链，也不是去做 SQL、POC、备份审计或工作台页面。"
        "C 的核心任务，是基于当前已经跑通的 Pikachu 8765 交付基线，提升 js_audit 的可读性、定位能力、误报控制和前端证据质量，"
        "让 frontend_audit 和 blackbox_pentest 中的前端结果更像真正的客户端安全审计结论。",
        [
            "项目目标不局限单一靶场。",
            "当前交付与验收默认基线是 Pikachu:8765。",
            "当前默认靶场是 Docker Pikachu，地址是 http://127.0.0.1:8765/。",
            "当前 frontend_audit 的默认模块只有 recon 和 js_audit。",
            "当前 js_audit 已经能抓页面脚本并做轻量规则分析，但结果偏多、偏散，仍需人工整理可读性。",
            "当前项目目标是授权靶场验证，不是前端攻击器，也不是大规模静态分析平台。",
        ],
        [
            "python_mvp/src/ai_security_agent/modules/js_audit.py",
            "python_mvp/src/ai_security_agent/modules/common.py",
            "python_mvp/src/ai_security_agent/schemas.py",
            "python_mvp/src/ai_security_agent/api/static/index.html",
            "python_mvp/tests/test_modules.py、test_agent_runtime.py",
            "python_mvp/skills/js_audit.yaml",
        ],
        [
            [r"python_mvp/src/ai_security_agent/modules/js_audit.py", "前端脚本抓取、规则检测、结果去重、证据表达"],
            [r"python_mvp/skills/js_audit.yaml", "维护 JS skill 的说明、触发词、检查清单"],
            [r"python_mvp/tests/test_modules.py", "补充与 js_audit 行为一致的测试"],
        ],
        [
            "不要修改 schemas.py 的字段结构。",
            "不要重写 api/app.py、api/service.py、runtime.py、report.py。",
            "不要把 js_audit 改成公网抓取器、浏览器攻击器或 payload 生成器。",
            "不要为了减少数量把所有高价值 finding 都删掉。",
        ],
        [
            (
                "阶段 0：理解边界",
                [
                    "确认目标是 Pikachu 8765 本地靶场，不是公网网站。",
                    "确认输出目标是更真实、更聚焦的前端审计结论，而不是更多条结果。",
                    "确认所有前端 finding 仍然必须适配当前统一报告结构。",
                ],
                "阶段 0 通过标准：能解释当前 js_audit 的问题不是没有功能，而是结果还不够适合交付和答辩展示。",
            ),
            (
                "阶段 1：优化脚本抓取与证据质量",
                [
                    "让结果里明确区分外部脚本、内联脚本、第三方库。",
                    "为每条 finding 补充更清楚的位置依据，例如脚本 URL、内联脚本编号、匹配模式。",
                    "如果公共库触发大量重复结果，要减少低价值噪音。",
                ],
                "阶段 1 通过标准：前端 finding 更像可解释的审计结论。",
            ),
            (
                "阶段 2：去重与误报控制",
                [
                    "对重复 sink 做合理聚合，避免刷满报告。",
                    "优先保留包含 eval、innerHTML、公开 API 路径、疑似密钥的高价值结果。",
                    "控制 findings 数量，优先保留 3 到 6 条适合展示的结论。",
                ],
                "阶段 2 通过标准：frontend_audit 结果更少但更有价值、更适合展示。",
            ),
        ],
        [
            r"cd D:\App\aster-main\aster-main",
            r"python -m unittest discover -s python_mvp/tests",
            r"用 D:\Python\python.exe 启动工作台，打开 http://127.0.0.1:8000/",
            r"目标填 http://127.0.0.1:8765/ ，Profile 选 frontend_audit",
            r"点击 Continue Until Pause，确认 recon 和 js_audit 执行完毕",
            r"生成报告并检查 findings 是否更聚焦、数量合理、定位清楚",
        ],
        [
            "frontend_audit 跑完后直接完成，不需要高风险确认。",
            "recon 能抓到 Pikachu 首页标题、脚本数量和基础入口信息。",
            "js_audit 的结果不再是一长串重复 sink，而是几条更聚焦的前端审计结论。",
        ],
        [
            "更新后的 js_audit.py",
            "同步过的 js_audit.yaml",
            "配套测试更新",
            "frontend_audit 完成页面截图和报告截图",
        ],
        "我没有改 Agent 主链和统一接口，我主要优化了前端 JS 审计结果的质量，让结果更聚焦、更可解释、更适合交付和答辩展示。",
    )


def build_06() -> None:
    build_member_doc(
        "06_D_task_handoff.docx",
        "成员 D 详细任务说明",
        "Recon、Backup 与靶场环境维护 / 基于当前 Pikachu 8765 交付基线",
        "D 的职责不是改 SQL、POC，也不是改前端 JS 审计。D 的核心任务，是把 recon、backup_audit、"
        "Pikachu 靶场维护、交付截图与环境说明做扎实，让项目的运行环境稳定、基础发现可重复复现、交付材料完整。",
        [
            "项目目标不局限单一靶场。",
            "当前交付与验收默认基线是 Pikachu:8765。",
            "当前默认靶场是 Docker Pikachu，地址是 http://127.0.0.1:8765/。",
            "当前工作台地址是 http://127.0.0.1:8000/。",
            "当前 Pikachu 容器偶尔会出现 8765 空响应，需要手工删除 apache pid 后重启 apache2。",
            "当前 recon 和 backup_audit 已经能工作，但还需要更稳定的靶场维护和更适合交付的证据说明。",
        ],
        [
            "python_mvp/src/ai_security_agent/modules/recon.py",
            "python_mvp/src/ai_security_agent/modules/backup_audit.py",
            "python_mvp/src/ai_security_agent/modules/common.py",
            "python_mvp/scripts/start_pikachu.ps1",
            "python_mvp/scripts/start_workbench.ps1",
            "docs/交付/02_交付说明_环境与部署.docx",
        ],
        [
            [r"python_mvp/src/ai_security_agent/modules/recon.py", "首页结构、标题、链接和脚本数量记录的稳定性"],
            [r"python_mvp/src/ai_security_agent/modules/backup_audit.py", "备份路径字典、结果可读性、敏感文件证据说明"],
            [r"python_mvp/src/ai_security_agent/modules/common.py", "必要时调整安全抓取超时和白名单逻辑"],
            [r"python_mvp/scripts/start_pikachu.ps1", "Pikachu 靶场启动脚本维护"],
            [r"python_mvp/scripts/start_workbench.ps1", "Workbench 启动脚本维护"],
            [r"docs/交付/02_交付说明_环境与部署.docx", "环境说明、路径说明、截图与交付记录同步"],
        ],
        [
            "不要修改 schemas.py 的字段结构。",
            "不要重写 runtime、service、planner、llm 相关主链逻辑。",
            "不要把 recon 或 backup_audit 扩成公网扫描器或大规模爬虫。",
            "不要改变 sql_scan、poc_verify 的高风险门禁逻辑。",
        ],
        [
            (
                "阶段 0：稳定靶场环境",
                [
                    "确认 Docker 容器 quirky_heyrovsky 可以启动。",
                    "确认 http://127.0.0.1:8765/ 能返回 Pikachu 首页。",
                    "如果出现空响应，整理清楚如何恢复 apache2 并写入交付说明。",
                ],
                "阶段 0 通过标准：别人照着命令能把 Pikachu 靶场启动起来。",
            ),
            (
                "阶段 1：完善 recon",
                [
                    "让 recon 对首页标题、链接数、脚本数、表单数的记录更稳定。",
                    "如果 Pikachu 首页结构变化，确保 evidence 仍然有解释价值。",
                    "必要时补充更适合交付展示的日志文字。",
                ],
                "阶段 1 通过标准：recon 的 finding 能稳定反映 Pikachu 首页的基本结构。",
            ),
            (
                "阶段 2：完善 backup_audit",
                [
                    "检查当前备份路径字典是否适合 Pikachu 环境。",
                    "如果发现 .git/config、备份包、配置文件暴露，要把 evidence 写得更适合答辩展示。",
                    "如果某些路径长期无价值，可以调整优先级，但不要破坏当前只读白名单探测模式。",
                ],
                "阶段 2 通过标准：backup_audit 报告结果更像交付级检查项，而不是一串生硬路径日志。",
            ),
            (
                "阶段 3：环境与交付证据整理",
                [
                    "补齐 Pikachu 首页截图、工作台截图、报告截图。",
                    "同步更新交付说明中的路径、端口、启动命令和已知问题。",
                    "保证别人按文档能复现实验环境。",
                ],
                "阶段 3 通过标准：交付说明、截图和环境步骤完整，不需要组长口头补充。",
            ),
        ],
        [
            r"cd D:\App\aster-main\aster-main",
            r"docker start quirky_heyrovsky",
            r"docker exec quirky_heyrovsky sh -lc ""rm -f /var/run/apache2/apache2.pid && apache2ctl start""",
            r"打开 http://127.0.0.1:8765/ 确认 Pikachu 首页可访问",
            r"打开 http://127.0.0.1:8000/ ，创建 frontend_audit 或 blackbox_pentest 任务",
            r"检查 recon 和 backup_audit 的结果是否稳定、可解释",
        ],
        [
            "Pikachu 靶场能稳定打开，不再出现端口有映射但浏览器空响应的情况。",
            "recon 能稳定记录标题、链接、脚本和入口信息。",
            "backup_audit 能输出少量高价值结果，例如 .git/config 暴露等。",
            "交付说明中的环境路径、脚本、命令和截图与当前代码一致。",
        ],
        [
            "更新后的 recon.py",
            "更新后的 backup_audit.py",
            "必要时更新后的 common.py",
            "稳定可用的 start_pikachu.ps1 和 start_workbench.ps1",
            "Pikachu 首页截图、工作台首页截图、recon/backup_audit 结果截图",
            "同步过的交付说明文档",
        ],
        "我主要负责把环境和基础发现打磨稳定。现在 Pikachu 8765 的启动链路更清楚，recon 和 backup_audit 的输出更适合交付说明和答辩展示，别人照着文档可以复现环境并看到基础结果。",
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_01()
    build_02()
    build_03()
    build_04()
    build_05()
    build_06()
    print("Generated delivery docs in", OUT_DIR)


if __name__ == "__main__":
    main()
